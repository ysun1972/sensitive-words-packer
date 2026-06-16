"""
多格式文件处理器
支持：
  - .txt / .md / .log / .csv（直接读写）
  - .docx（python-docx，保留段落 + 表格）
  - .pdf（pypdf 提取 + reportlab 重写 PDF + 同文本生成 docx 双产物）

⚠️ PDF 处理说明：
  PDF 内部是二进制流（字体子集、坐标），无法"原地替换"敏感词。
  本工具采用"提取 → 脱敏 → 重建"策略：
    1. 读：从 PDF 提取纯文本（pypdf）
    2. 写（PDF 模式 C）：
       - 主输出：reportlab 重建的简化版 PDF（保留文件名/扩展名）
       - 副输出：同名 .docx（用同一脱敏文本，便于用户后续编辑）
    限制：版式（页眉页脚、表格、图表、字体）会简化为纯文本样式；
         扫描版 PDF（图片型）无文本层，read() 返回空字符串。
"""
from __future__ import annotations

import shutil
from pathlib import Path
from typing import Protocol, runtime_checkable


@runtime_checkable
class FileHandler(Protocol):
    """文件处理器接口"""
    extensions: tuple[str, ...]
    # 是否支持双产物（PDF 默认 True，会额外输出同名 .docx）
    supports_dual_output: bool

    def read(self, path: Path) -> str: ...
    def write(self, path: Path, text: str) -> list[Path]:
        """写入文件，返回所有产物路径列表"""
        ...


class TextHandler:
    """处理纯文本文件（.txt / .md / .log / .csv）"""
    extensions = (".txt", ".md", ".log", ".csv")
    supports_dual_output = False

    def read(self, path: Path) -> str:
        # 尝试 utf-8，失败回退 gbk（兼容 Windows GBK 编码）
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="gbk", errors="ignore")

    def write(self, path: Path, text: str) -> list[Path]:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(text, encoding="utf-8")
        return [path]


class DocxHandler:
    """处理 .docx 文件（通过 python-docx）"""
    extensions = (".docx",)
    supports_dual_output = False

    def __init__(self):
        try:
            import docx  # noqa: F401
            self._available = True
        except ImportError:
            self._available = False

    def read(self, path: Path) -> str:
        if not self._available:
            raise RuntimeError("需要安装 python-docx: pip install python-docx")
        from docx import Document
        doc = Document(str(path))
        # 段落 + 表格，按出现顺序拼接
        parts: list[str] = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    def write(self, path: Path, text: str) -> list[Path]:
        if not self._available:
            raise RuntimeError("需要安装 python-docx: pip install python-docx")
        from docx import Document
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(path))
        return [path]


class PdfHandler:
    """
    处理 .pdf 文件（双产物模式 C）

    读：pypdf 提取每页文本
    写：
      - 主产物：reportlab 重建的简化版 PDF（保留文件名/扩展名）
      - 副产物：同名 .docx（用同一脱敏文本，便于用户编辑）

    使用依赖：
      - pypdf（读取）
      - reportlab（PDF 写入，内置 STSong-Light 字体支持中文）
      - python-docx（docx 副产物写入）
    """
    extensions = (".pdf",)
    supports_dual_output = True

    def __init__(self):
        self._available = False
        self._deps_error: str = ""
        try:
            import pypdf  # noqa: F401
        except ImportError:
            self._deps_error += "缺少 pypdf，请 pip install pypdf\n"
        try:
            import reportlab  # noqa: F401
        except ImportError:
            self._deps_error += "缺少 reportlab，请 pip install reportlab\n"
        try:
            import docx  # noqa: F401
        except ImportError:
            self._deps_error += "缺少 python-docx（用于 docx 副产物），请 pip install python-docx\n"
        if not self._deps_error:
            self._available = True

    def read(self, path: Path) -> str:
        if not self._available:
            raise RuntimeError(self._deps_error)
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        parts: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception as e:
                text = f"[第 {i+1} 页提取失败: {e}]"
            if text.strip():
                parts.append(text)
        return "\n".join(parts)

    def write(self, path: Path, text: str) -> list[Path]:
        """
        写双产物：
          - path (e.g. foo.pdf) → 主 PDF 产物
          - path.with_suffix(".docx") (e.g. foo.docx) → 副 docx 产物
        """
        if not self._available:
            raise RuntimeError(self._deps_error)
        path.parent.mkdir(parents=True, exist_ok=True)

        products: list[Path] = []

        # 产物 1: PDF（reportlab 重建）
        products.append(self._write_pdf(path, text))

        # 产物 2: 同名 .docx（便于用户编辑 / 排版）
        docx_path = path.with_suffix(".docx")
        products.append(self._write_docx(docx_path, text))

        return products

    def _write_pdf(self, path: Path, text: str) -> Path:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        # 注册内置中文字体（reportlab 自带 CID 字体，无需外部字体文件）
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
        except Exception:
            font_name = "Helvetica"

        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
            title=path.stem,
        )
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            "BodyCN",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=14,
        )
        story = []
        for line in text.split("\n"):
            if not line.strip():
                story.append(Spacer(1, 6))
                continue
            safe = (
                line.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
            )
            story.append(Paragraph(safe, body_style))
        doc.build(story)
        return path

    def _write_docx(self, path: Path, text: str) -> Path:
        from docx import Document
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(path))
        return path


# 注册表
_HANDLERS: list[FileHandler] = [TextHandler(), DocxHandler(), PdfHandler()]
_HANDLER_MAP: dict[str, FileHandler] = {}
for h in _HANDLERS:
    for ext in h.extensions:
        _HANDLER_MAP[ext.lower()] = h


def get_handler(path: Path) -> FileHandler | None:
    """根据文件后缀返回对应 handler，不支持则返回 None"""
    return _HANDLER_MAP.get(path.suffix.lower())


def supported_extensions() -> list[str]:
    """返回所有支持的文件后缀"""
    return list(_HANDLER_MAP.keys())


def copy_file(path: Path, output_dir: Path) -> list[Path]:
    """不支持的格式直接复制（保路径）"""
    rel = path.name
    dest = output_dir / rel
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, dest)
    return [dest]
